from __future__ import annotations

import re
import shutil
from pathlib import Path

import docx
import pymupdf

from doc_to_abstract.exceptions import ConfigError


def read_template(template_path: str) -> str:
    """Read a template file (.tex or .docx) and return its text content."""
    path = Path(template_path)
    if not path.exists():
        raise ConfigError(f"Template file not found: {template_path}")

    suffix = path.suffix.lower()
    if suffix == ".tex":
        return path.read_text(encoding="utf-8")
    elif suffix == ".docx":
        return _read_docx(template_path)
    elif suffix == ".pdf":
        return _read_pdf(template_path)
    else:
        raise ConfigError(f"Unsupported template format: {suffix} (use .tex, .docx, or .pdf)")


def fill_template(template_path: str, abstract_text: str, output_path: str) -> None:
    """Copy the template and insert the abstract text into it."""
    path = Path(template_path)
    suffix = path.suffix.lower()

    if suffix == ".tex":
        _fill_tex(template_path, abstract_text, output_path)
    elif suffix == ".docx":
        _fill_docx(template_path, abstract_text, output_path)
    else:
        raise ConfigError(f"Cannot fill template format: {suffix} (use .tex or .docx)")


def _read_docx(docx_path: str) -> str:
    """Extract all text from a .docx file."""
    doc = docx.Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _read_pdf(pdf_path: str) -> str:
    """Extract all text from a PDF template."""
    doc = pymupdf.open(pdf_path)
    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def _fill_tex(template_path: str, abstract_text: str, output_path: str) -> None:
    """Insert abstract into a LaTeX template's abstract environment."""
    content = Path(template_path).read_text(encoding="utf-8")

    # Replace the content between \begin{abstract} and \end{abstract}
    pattern = r"(\\begin\{abstract\})(.*?)(\\end\{abstract\})"
    replacement = rf"\1\n{abstract_text}\n\3"
    new_content, count = re.subn(pattern, replacement, content, flags=re.DOTALL)

    if count == 0:
        # No abstract environment found; append before \end{document}
        end_doc = r"\end{document}"
        if end_doc in new_content:
            insert = f"\n\\begin{{abstract}}\n{abstract_text}\n\\end{{abstract}}\n\n"
            new_content = new_content.replace(end_doc, insert + end_doc)
        else:
            # Just append
            new_content += f"\n\\begin{{abstract}}\n{abstract_text}\n\\end{{abstract}}\n"

    Path(output_path).write_text(new_content, encoding="utf-8")


def _fill_docx(template_path: str, abstract_text: str, output_path: str) -> None:
    """Insert abstract into a .docx template.

    Looks for a paragraph containing 'abstract' (case-insensitive) and inserts
    the generated text after it. If not found, appends at the end.
    """
    shutil.copy2(template_path, output_path)
    doc = docx.Document(output_path)

    # Find the abstract section
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        if "abstract" in text_lower and len(para.text.strip()) < 50:
            insert_index = i + 1
            break

    if insert_index is not None:
        # Remove existing placeholder text between abstract heading and next section
        # (paragraphs that look like placeholder/dummy text)
        while insert_index < len(doc.paragraphs):
            next_text = doc.paragraphs[insert_index].text.strip()
            # Stop if we hit another section heading or significant content marker
            if not next_text:
                insert_index += 1
                continue
            # Assume short headings or keyword-like lines are section boundaries
            if next_text.lower() in ("keywords", "keyword", "introduction", "1. introduction"):
                break
            # Remove placeholder paragraph by clearing it
            doc.paragraphs[insert_index].text = ""
            insert_index += 1

        # Insert abstract text after the heading
        target = doc.paragraphs[insert_index - 1] if insert_index > 0 else doc.paragraphs[0]
        new_para = doc.add_paragraph(abstract_text)
        # Move the new paragraph after the target
        target._element.addnext(new_para._element)
    else:
        # No abstract heading found; just append
        doc.add_paragraph(abstract_text)

    doc.save(output_path)
